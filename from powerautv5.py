

from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from docx import Document
from selenium.webdriver.common.by import By
import openpyxl
import time

# Initialize Chrome with options, if needed
options = Options()
# options.add_argument('--headless')  # Uncomment for headless mode

# Initialize the driver
driver = webdriver.Chrome(options=options)

urlinf = "https://finance.yahoo.com/quote/IBM/financials"
urlinf1 = "https://www.asx.com.au/markets/company/BHP"
urlinf2 = "https://www.asx.com.au/markets/company/ACR"


# Navigate to the Yahoo Finance IBM financials page
driver.get(urlinf1)

# Allow some time for the page to load
time.sleep(5)

# Extract the financial data
financial_data1 = driver.find_element(By.TAG_NAME, "body").text

# Navigate to the Yahoo Finance IBM financials page
driver.get(urlinf2)

# Allow some time for the page to load
time.sleep(5)

# Extract the financial data
financial_data2 = driver.find_element(By.TAG_NAME, "body").text

# Close the browser
driver.quit()



# Create a new Word document and add the extracted text
doc = Document()
doc.add_paragraph(financial_data1)
doc.add_paragraph(financial_data2)
doc.save('financial_data.docx')

# Create a new Excel workbook and add the extracted text
wb = openpyxl.Workbook()
ws = wb.active
ws.title = "Financial Data"

# Define function to create a new sheet and write data
def write_to_new_sheet(sheet_name, data):
  # Create a new worksheet
  new_sheet = wb.create_sheet(sheet_name)
  # Write the data row by row
  for row_num, row_data in enumerate(data.split('\n'), 1):
    new_sheet.cell(row=row_num, column=1, value=row_data)

# Write data to "financial_data" sheet
write_to_new_sheet("Financial_data1", financial_data1)

# Write data to "financial_data2" sheet (assuming financial_data2 is defined)
write_to_new_sheet("Financial_data2", financial_data2)


# Save the Excel workbook
xlname = 'financial_data_3.xlsx'
wb.save(xlname)


print(f"Text has been copied and saved to 'financial_data.docx' and '{xlname}'")

# Note: you may need to install openpyxl using: pip install selenium python-docx openpyxl


