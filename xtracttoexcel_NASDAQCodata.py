

from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from docx import Document
from selenium.webdriver.common.by import By
import openpyxl
import time
import random

# Initialize Chrome with options, if needed
options = Options()
# options.add_argument('--headless')  # Uncomment for headless mode

# Initialize the driver
driver = webdriver.Chrome(options=options) # You can also use other web drivers like Firefox or Edge


# Example list of ASX codes
NASDAQ_symbols = ['BAC', 'CVX', 'NFLX', 'AZN']  # Add more NASDAQ Symbols as needed

# Create a new Word document
doc = Document()

# Create a new Excel workbook and get the active sheet
wb = openpyxl.Workbook()
ws = wb.active
ws.title = "NASDAQ_FinancialData"


# Function to write data to a new column in the Excel sheet
def write_to_new_column(sheet, data):
  # Find the next empty column
  if sheet.max_column == 1 and sheet.cell(1, 1).value is None:
    new_column = 1
  else:
    new_column = sheet.max_column + 1

  # Write the data row by row to the next empty column
  for row_num, row_data in enumerate(data.split('\n'), 1):
    sheet.cell(row=row_num, column=new_column, value=row_data)


# Iterate through each ASX code, extract financial data, and write to the document and workbook
for symbol in NASDAQ_symbols:
  url = f"https://www.nasdaq.com/market-activity/stocks/{symbol}"

  # Navigate to the web page
  driver.get(url)
  time.sleep(5)  # Allow time for the page to load
  # Generate a random sleep time between 5 and 20 seconds
  sleep_time = random.uniform(5, 15)
  #sleep_time = random.randint(5, 20)
  time.sleep(sleep_time)  # Allow time for the page to load

  # Extract financial data
  financial_data = driver.find_element(By.TAG_NAME, "body").text

  # Add extracted text to Word document
  doc.add_paragraph(f"NASDAQ Symbol: {symbol}")
  doc.add_paragraph(financial_data)

  # Write extracted text to the Excel sheet
  write_to_new_column(ws, f"NASDAQ Symbol: {symbol}\n{financial_data}")

# Close the browser
driver.quit()

# Save the Word document
doc.save('financial_data.docx')

# Save the Excel workbook
wb.save('NASDAQ_FinancialData.xlsx')

print("Text has been copied and saved to 'financial_data.docx' and 'NASDAQ_FinancialData.xlsx'")
# Note: you may need to install openpyxl using: pip install selenium python-docx openpyxl


