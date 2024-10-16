

from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from docx import Document
from selenium.webdriver.common.by import By
import openpyxl
import time
import random

# Initialize Chrome with options, if needed
options = Options()
# options.add_argument('--headless')  # Uncomment for headless mode

# Initialize the driver
driver = webdriver.Chrome(options=options) # You can also use other web drivers like Firefox or Edge


# Example list of ASX codes
asx_codes = ['BHP', 'CBA', 'CSL', 'NAB', 'WBC', 'ANZ', 'WES', 'MQG', 'FMG', 'NEM', 'GMG', 'SQ2', 'WDS', 'RMD', 'RIO', 'TLS', 'WOW', 'TCL', 'WTC', 'ALL', 'QBE', 'REA', 'STO', 'NWS', 'NWSLV', 'COL', 'AMC', 'COH', 'JHX', 'BXB', 'XRO', 'ORG', 'REH', 'URW', 'FPH', 'S32', 'SCG', 'CPU', 'SVW', 'NST', 'IAG', 'MEZ', 'PME', 'CAR', 'LNW', 'MIN', 'SHL', 'SOL', 'TLC', 'ASX', 'RHC', 'NXT', 'APA', 'SGP', 'AIA', 'MPL', 'PLS', 'QAN', 'TWE', 'EDV', 'AFI', 'BSL', 'ALU', 'ORI', 'TPG', 'VCX', 'MCY', 'CSC', 'SEK', 'IFT', 'YAL', 'GQG', 'ALD', 'GPT', 'MGR', 'WOR', 'EVN', 'DXS', 'ALX', 'CEN', 'ALQ', 'AGL', 'JBH', 'SPK', 'AZJ', 'ARG', 'QUB', 'WHC', 'BEN', 'CWY', 'SDF']  # Add more ASX codes as needed
#List1 'BHP', 'CBA', 'CSL', 'NAB', 'WBC', 'ANZ', 'WES', 'MQG', 'FMG', 'NEM', 'GMG', 'SQ2', 'WDS', 'RMD', 'RIO', 'TLS', 'WOW', 'TCL', 'WTC', 'ALL', 'QBE', 'REA', 'STO', 'NWS', 'NWSLV', 'COL', 'AMC', 'COH', 'JHX', 'BXB', 'XRO', 'ORG', 'REH', 'URW', 'FPH', 'S32', 'SCG', 'CPU', 'SVW', 'NST', 'IAG', 'MEZ', 'PME', 'CAR', 'LNW', 'MIN', 'SHL', 'SOL', 'TLC', 'ASX', 'RHC', 'NXT', 'APA', 'SGP', 'AIA', 'MPL', 'PLS', 'QAN', 'TWE', 'EDV', 'AFI', 'BSL', 'ALU', 'ORI', 'TPG', 'VCX', 'MCY', 'CSC', 'SEK', 'IFT', 'YAL', 'GQG', 'ALD', 'GPT', 'MGR', 'WOR', 'EVN', 'DXS', 'ALX', 'CEN', 'ALQ', 'AGL', 'JBH', 'SPK', 'AZJ', 'ARG', 'QUB', 'WHC', 'BEN', 'CWY', 'SDF'
#List2 'LTM', 'TNE', 'LYC', 'CHC', 'EBO', 'NXG', 'IPL', 'HVN', 'TLX', 'VUK', 'NWL', 'VEA', 'ATM', 'A2M', 'AWC', 'IGO', 'PMV', 'CGF', 'CSR', 'FLT', 'PDN', 'IEL', 'MTS', 'BKW', 'BRG', 'NHC', 'SFR', 'BOQ', 'LLC', 'RWC', 'ANN', 'BPT', 'HUB', 'NIC', 'AUB', 'NHF', 'DMP', 'WEB', 'LOV', 'CIA', 'MGF', 'PRU', 'ARB', 'VNT', '360', 'NSR', 'DOW', 'CNU', 'SMR', 'BFL', 'SUL', 'ILU', 'AMP', 'ORA', 'PNI', 'HMC', 'APE', 'LTR', 'PXA', 'BWP', 'RGN', 'NEU', 'HDN', 'CLW', 'DEG', 'PPT', 'EMR', 'DRR', 'MAQ', 'BGL', 'NEC', 'PSI', 'GNE', 'MXT', 'FBU', 'RMS', 'SNZ', 'ABC', 'MFF', 'CDA', 'MP1', 'SGM', 'TUA', 'CIP', 'GMD', 'CRN', 'INA', 'GNC', 'LSF', 'CTD', 'CQR', 'SIG', 'DHG', 'ZIM', 'NUF', 'EVT', 'GOZ', 'DDR', 'BAP', 'GOR', 'AIZ', 'CMM', 'CU6', 'BOE', 'JLG', 'PNV', 'WLE', 'ASK', 'WAM', 'RDX', 'IPH', 'WPR', 'MAC', 'RED', 'REDNB', 'LIC', 'ZIP', 'SSR', 'GUD', 'IRE', 'MFG', 'IFL', 'JDO', 'WAF', 'CNI', 'DBI', 'DYL', 'AD8', 'SGR', 'MGH', 'TAH', 'SDR', 'BKI', 'KLS', 'NWH', 'KAR', 'BGA', 'RRL', 'ARF', 'ING', 'DTL', 'FRW', 'SLX', 'AUI', 'ADT', 'MND', 'ELD', 'APM', 'REG', 'MSB', 'HLI', 'MAD', 'LFS', 'NCK', 'TPW', 'MMS', 'RSG', 'OCL', 'IMD', 'LFG', 'PWH', 'WGX', 'AX1', 'DUI', 'DRO', 'CMW', 'SIQ', 'CKF', 'ABB', 'SKC', 'HLS', 'MRM', 'ABG', 'SGF', 'JIN', 'NGI', 'CCP', 'GEM', 'GTK', 'SNL', 'VSL', 'WA1', 'BGP', 'NXL', 'UOS', 'DXI', 'PRN', 'SPR', 'RPL', 'PGF', 'VUL', 'PL8', 'CQE', 'NAN', 'HSN', 'ERA', 'HGH', 'CTT', 'AAC', 'FPR', 'ASB', 'EQT', 'A4N', 'MAF', 'WGB', 'PFP', 'CIN', 'RFF', 'DJW', 'CBO'

# Create a new Word document
#doc = Document()

# Create a new Excel workbook and get the active sheet
wb = openpyxl.Workbook()
ws = wb.active
ws.title = "Financial_Data"


# Function to write data to a new column in the Excel sheet
def write_to_new_column(sheet, data):
  # Find the next empty column
  if sheet.max_column == 1 and sheet.cell(1, 1).value is None:
    new_column = 1
  else:
    new_column = sheet.max_column + 1

  # Write the data row by row to the next empty column
  for row_num, row_data in enumerate(data.split('\n'), 1):
    sheet.cell(row=row_num, column=new_column, value=row_data)


# Iterate through each ASX code, extract financial data, and write to the document and workbook
for code in asx_codes:
  url = f"https://www.asx.com.au/markets/company/{code}"

  # Navigate to the web page
  driver.get(url)
  time.sleep(5)  # Allow time for the page to load
  # Generate a random sleep time between 5 and 20 seconds
  sleep_time = random.uniform(5, 15)
  #sleep_time = random.randint(5, 20)
  time.sleep(sleep_time)  # Allow time for the page to load

  # Extract financial data
  financial_data = driver.find_element(By.TAG_NAME, "body").text

  # Add extracted text to Word document
  #doc.add_paragraph(f"ASX Code: {code}")
  #doc.add_paragraph(financial_data)

  # Write extracted text to the Excel sheet
  write_to_new_column(ws, f"ASX Code: {code}\n{financial_data}")

# Close the browser
driver.quit()

# Save the Word document
#doc.save('financial_data.docx')

# Save the Excel workbook
wb.save('financial_data.xlsx')

print("Text has been copied and saved to 'financial_data.docx' and 'financial_data.xlsx'")
# Note: you may need to install openpyxl using: pip install selenium python-docx openpyxl


