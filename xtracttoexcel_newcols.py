

from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from docx import Document
from selenium.webdriver.common.by import By
import openpyxl
import time

# Initialize Chrome with options, if needed
options = Options()
# options.add_argument('--headless')  # Uncomment for headless mode

# Initialize the driver
driver = webdriver.Chrome(options=options)

urlinf1 = "https://www.asx.com.au/markets/company/BHP"
urlinf2 = "https://www.asx.com.au/markets/company/ACR"


# Navigate to the Yahoo Finance IBM financials page
driver.get(urlinf1)

# Allow some time for the page to load
time.sleep(5)

# Extract the financial data
financial_data1 = driver.find_element(By.TAG_NAME, "body").text

# Navigate to the Yahoo Finance IBM financials page
driver.get(urlinf2)

# Allow some time for the page to load
time.sleep(5)

# Extract the financial data
financial_data2 = driver.find_element(By.TAG_NAME, "body").text

# Close the browser
driver.quit()



# Create a new Word document and add the extracted text
doc = Document()
doc.add_paragraph(financial_data1)
doc.add_paragraph(financial_data2)
doc.save('financial_data.docx')

# Create a new Excel workbook and add the extracted text
wb = openpyxl.Workbook()
ws = wb.active
ws.title = "Financial_Data"



def write_to_new_column(sheet_name, data):
  sheet = wb[sheet_name]

  # Find the next empty column
  if sheet.max_column == 1 and sheet.cell(1, 1).value is None:
    new_column = 1
  else:
    new_column = sheet.max_column + 1

  # Write the data row by row to the next empty column
  for row_num, row_data in enumerate(data.split('\n'), 1):
    sheet.cell(row=row_num, column=new_column, value=row_data)


# Assuming financial_data1 and financial_data2 are defined
# Call the function to write data to the columns within the existing sheet
write_to_new_column("Financial_Data", financial_data1)
write_to_new_column("Financial_Data", financial_data2)

# Save the workbook with changes
wb.save('workbook.xlsx')




# Save the Excel workbook
xlname = 'financial_data_4.xlsx'
wb.save(xlname)


print(f"Text has been copied and saved to 'financial_data.docx' and '{xlname}'")

# Note: you may need to install openpyxl using: pip install selenium python-docx openpyxl


